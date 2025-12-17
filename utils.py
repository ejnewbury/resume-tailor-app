# utils.py - Utility functions for resume tailoring (no Streamlit dependencies)

import os
import pdfplumber
import requests
from bs4 import BeautifulSoup
import spacy
from collections import Counter
from sentence_transformers import SentenceTransformer, util
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
from openai import OpenAI
import urllib.parse
from jinja2 import Template
import torch
import pandas as pd

# ------------------- LLM CLIENT CONFIGURATION -------------------
def get_llm_client(llm_config=None):
    """Create OpenAI client with environment variable configuration for maximum compatibility."""
    if not llm_config:
        llm_config = {
            'type': 'local',
            'model': 'llama2',
            'server': 'http://localhost:11434'
        }

    llm_type = llm_config.get('type', 'local')

    # Use environment variables with sensible defaults for different providers
    if llm_type == 'openai':
        base_url = os.getenv("LLM_BASE_URL", "https://api.openai.com/v1")
        api_key = os.getenv("LLM_API_KEY", llm_config.get('api_key', ''))
    else:  # local LLM (LM Studio, Ollama, etc.)
        # Default to LM Studio format, but allow override
        if 'localhost:11434' in llm_config.get('server', ''):
            # Ollama - convert to OpenAI-compatible format
            base_url = os.getenv("LLM_BASE_URL", "http://localhost:11434/v1")
        else:
            # LM Studio or other OpenAI-compatible servers
            base_url = os.getenv("LLM_BASE_URL", "http://localhost:1234/v1")

        api_key = os.getenv("LLM_API_KEY", "not-needed")  # Local servers don't need real API keys

    return OpenAI(
        base_url=base_url,
        api_key=api_key
    )

# ------------------- MODEL LOADING -------------------
def load_models(device_override=None):
    """Load models with device selection and error handling."""
    try:
        # Use override if provided, otherwise default to auto (CUDA if available)
        device = device_override if device_override else ('cuda' if torch.cuda.is_available() else 'cpu')

        # Load sentence transformer
        embedding_model = SentenceTransformer('all-MiniLM-L6-v2', device=device)

        # Test the model with a small input to catch issues early
        test_text = "test"
        embedding_model.encode(test_text, convert_to_tensor=False)

        # Load spaCy
        nlp_model = spacy.load("en_core_web_sm")

        print(f"Models loaded successfully on device: {device}")
        return embedding_model, nlp_model

    except Exception as e:
        error_msg = str(e).lower()
        if 'cuda' in error_msg or 'gpu' in error_msg or 'kernel' in error_msg:
            print("CUDA/GPU error detected. Switching to CPU mode...")
            try:
                # Force CPU mode
                embedding_model = SentenceTransformer('all-MiniLM-L6-v2', device='cpu')
                nlp_model = spacy.load("en_core_web_sm")

                print("Using CPU mode - performance may be slower but more stable")
                return embedding_model, nlp_model

            except Exception as e2:
                print(f"Critical error: Could not load models: {str(e2)}")
                raise RuntimeError("Failed to load required models")
        else:
            print(f"Model loading error: {str(e)}")
            raise RuntimeError("Failed to load required models")

# ------------------- COVER LETTER TEMPLATES -------------------

COVER_LETTER_TEMPLATES = {
    'professional': Template("""
{{ opening }}

{{ body_intro }}

{{ achievements }}

{{ skills_match }}

{{ closing }}

Sincerely,
{{ candidate_name }}
"""),

    'creative': Template("""
{{ opening }}

{{ body_intro }}

{{ achievements }}

{{ skills_match }}

{{ closing }}

Best regards,
{{ candidate_name }}
"""),

    'concise': Template("""
{{ opening }}

{{ body_paragraph }}

{{ closing }}

Regards,
{{ candidate_name }}
""")
}

TONE_PROMPTS = {
    'professional': """Write in a formal, professional tone. Use industry-standard language and maintain a respectful, business-appropriate style.
Focus on qualifications, experience, and demonstrated results. Avoid casual language or overly creative expressions.
Structure: Clear opening stating the position, body paragraphs highlighting relevant experience with metrics, and a confident closing.""",

    'creative': """Write in an engaging, memorable tone that shows personality while remaining professional.
Use compelling storytelling to highlight achievements. Include a unique hook in the opening that captures attention.
Structure: Start with an engaging hook, tell your story through achievements, and end with a memorable close.""",

    'concise': """Write in a direct, no-nonsense style. Get to the point quickly and focus on the most relevant qualifications.
Use bullet points where appropriate for clarity. Keep it under 300 words total.
Structure: State your interest and qualifications in the first paragraph, provide key evidence in the second, and close confidently."""
}

# ------------------- SKILL ANALYSIS -------------------

SKILL_ALIASES = {
    'javascript': ['js', 'javascript', 'ecmascript'],
    'typescript': ['ts', 'typescript'],
    'python': ['python', 'py'],
    'machine learning': ['ml', 'machine learning', 'machine-learning'],
    'artificial intelligence': ['ai', 'artificial intelligence'],
    'amazon web services': ['aws', 'amazon web services'],
    'google cloud platform': ['gcp', 'google cloud', 'google cloud platform'],
    'microsoft azure': ['azure', 'microsoft azure'],
    'continuous integration': ['ci', 'continuous integration'],
    'continuous deployment': ['cd', 'continuous deployment'],
    'ci/cd': ['ci/cd', 'cicd', 'ci cd'],
    'kubernetes': ['k8s', 'kubernetes'],
    'postgresql': ['postgres', 'postgresql'],
    'mongodb': ['mongo', 'mongodb'],
    'react': ['react', 'reactjs', 'react.js'],
    'angular': ['angular', 'angularjs'],
    'vue': ['vue', 'vuejs', 'vue.js'],
    'node.js': ['node', 'nodejs', 'node.js'],
    'sql': ['sql', 'structured query language'],
    'nosql': ['nosql', 'no-sql'],
    'api': ['api', 'apis', 'rest api', 'restful'],
    'agile': ['agile', 'scrum', 'kanban'],
    'project management': ['project management', 'pm', 'program management'],
    'data analysis': ['data analysis', 'data analytics', 'analytics'],
    'data science': ['data science', 'data scientist'],
    'deep learning': ['deep learning', 'dl', 'neural networks'],
    'natural language processing': ['nlp', 'natural language processing'],
    'computer vision': ['cv', 'computer vision', 'image processing'],
}

SKILL_CATEGORIES = {
    'programming': ['python', 'java', 'javascript', 'typescript', 'c++', 'c#', 'ruby', 'go', 'rust', 'php', 'swift', 'kotlin', 'scala', 'r'],
    'frameworks': ['react', 'angular', 'vue', 'django', 'flask', 'spring', 'express', 'rails', 'laravel', 'nextjs', 'fastapi'],
    'databases': ['sql', 'mysql', 'postgresql', 'mongodb', 'redis', 'elasticsearch', 'cassandra', 'dynamodb', 'oracle', 'sqlite'],
    'cloud': ['aws', 'azure', 'gcp', 'docker', 'kubernetes', 'terraform', 'jenkins', 'circleci', 'github actions'],
    'data_science': ['machine learning', 'deep learning', 'tensorflow', 'pytorch', 'pandas', 'numpy', 'scikit-learn', 'spark', 'hadoop'],
    'soft_skills': ['leadership', 'communication', 'teamwork', 'problem solving', 'analytical', 'strategic', 'collaboration', 'mentoring'],
    'tools': ['git', 'jira', 'confluence', 'slack', 'figma', 'photoshop', 'excel', 'tableau', 'power bi'],
}

def extract_candidate_name(resume_text):
    """Extract candidate name from resume text using simple heuristics."""
    lines = resume_text.split('\n')[:10]  # Check first 10 lines

    for line in lines:
        line = line.strip()
        if line and len(line.split()) >= 2 and len(line) < 50:  # Reasonable name length
            # Skip lines that look like addresses, emails, or phone numbers
            if '@' in line or any(char.isdigit() for char in line[:15]) or line.startswith('http'):
                continue
            # Skip lines that are all caps (likely section headers)
            if line.isupper():
                continue
            return line.title()  # Title case for names

    return "Candidate Name"

def extract_company_info(jd_text):
    """Extract company name from job description."""
    # Look for common company indicators
    indicators = ['at ', 'for ', 'with ', 'join ']
    lines = jd_text.split('\n')[:20]  # Check first 20 lines

    for line in lines:
        line = line.lower().strip()
        for indicator in indicators:
            if indicator in line:
                # Extract potential company name
                parts = line.split(indicator, 1)
                if len(parts) > 1:
                    company_part = parts[1].strip()
                    # Take first 1-3 words as company name
                    words = company_part.split()[:3]
                    if words:
                        return ' '.join(words).title()

    return "the company"

def generate_cover_letter(resume_text, tailored_resume, jd_text, tone, keywords):
    """Generate a personalized cover letter using OpenAI."""
    try:
        client = OpenAI(api_key=os.getenv('OPENAI_API_KEY'))

        candidate_name = extract_candidate_name(resume_text)
        company_name = extract_company_info(jd_text)

        # Extract key skills from keywords
        key_skills = keywords[:5] if isinstance(keywords, list) else []

        prompt = f"""Write a compelling cover letter for a software engineering position.

CANDIDATE: {candidate_name}
COMPANY: {company_name}
TONE: {tone}

RESUME SUMMARY:
{tailored_resume[:1000]}

JOB DESCRIPTION:
{jd_text[:1000]}

KEY SKILLS: {', '.join(key_skills)}

{TONE_PROMPTS[tone]}

Requirements:
- Keep it under 400 words
- Make it personalized and specific
- Highlight relevant experience and achievements
- Show enthusiasm for the role and company
- End with a call to action"""

        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}],
            max_tokens=800,
            temperature=0.7
        )

        return response.choices[0].message.content.strip()

    except Exception as e:
        print(f"Cover letter generation failed: {e}")
        return f"Dear Hiring Manager,\n\nI am excited to apply for this position at {company_name}. With my background in software development, I am confident I can contribute effectively to your team.\n\nPlease find my resume attached for your review.\n\nBest regards,\n{candidate_name}"

def create_cover_letter_docx(cover_letter_text, filename="Cover_Letter.docx"):
    """Create a professional cover letter DOCX file."""
    try:
        doc = Document()

        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Split cover letter into paragraphs
        paragraphs = [p.strip() for p in cover_letter_text.split('\n\n') if p.strip()]

        for i, para in enumerate(paragraphs):
            if i == 0:  # First paragraph - contact info placeholder
                p = doc.add_paragraph()
                p.add_run("Your Name\nYour Address\nCity, State, ZIP Code\nEmail | Phone | LinkedIn\n\n").font.size = Pt(10)
                p.add_run(para).font.size = Pt(11)
            elif i == len(paragraphs) - 1:  # Last paragraph - signature
                p = doc.add_paragraph()
                p.add_run(para).font.size = Pt(11)
                doc.add_paragraph("\nSincerely,")
                p = doc.add_paragraph()
                p.add_run("Your Name").font.size = Pt(12)
            else:
                p = doc.add_paragraph()
                p.add_run(para).font.size = Pt(11)

            # Add spacing between paragraphs
            if i < len(paragraphs) - 1:
                p.paragraph_format.space_after = Pt(6)

        # Save file
        output_path = os.path.join("outputs", filename)
        os.makedirs("outputs", exist_ok=True)
        doc.save(output_path)
        return output_path

    except Exception as e:
        print(f"Error creating cover letter DOCX: {e}")
        return None

def extract_skills(text, nlp_model):
    """Extract skills from text using spaCy NLP and pattern matching."""
    if not nlp_model:
        return set()

    doc = nlp_model(text.lower())
    skills = set()

    noun_phrases = [chunk.text.strip() for chunk in doc.noun_chunks]
    tokens = [token.lemma_ for token in doc if token.is_alpha and len(token.text) > 1]
    candidates = set(noun_phrases + tokens)

    all_known_skills = set()
    for category_skills in SKILL_CATEGORIES.values():
        all_known_skills.update(category_skills)
    for skill, aliases in SKILL_ALIASES.items():
        all_known_skills.add(skill)
        all_known_skills.update(aliases)

    for candidate in candidates:
        candidate_clean = candidate.strip().lower()
        if candidate_clean in all_known_skills:
            normalized = normalize_skill(candidate_clean)
            skills.add(normalized)
            continue
        for skill in all_known_skills:
            if skill in candidate_clean or candidate_clean in skill:
                normalized = normalize_skill(skill)
                skills.add(normalized)

    text_lower = text.lower()
    skill_patterns = [
        'experience with', 'proficient in', 'knowledge of', 'skilled in',
        'expertise in', 'familiar with', 'working knowledge of', 'hands-on experience',
        'strong background in', 'demonstrated ability'
    ]

    for pattern in skill_patterns:
        if pattern in text_lower:
            idx = text_lower.find(pattern)
            snippet = text_lower[idx:idx+100]
            snippet_doc = nlp_model(snippet)
            for token in snippet_doc:
                if token.lemma_ in all_known_skills:
                    skills.add(normalize_skill(token.lemma_))

    return skills

def normalize_skill(skill):
    """Normalize skill name to a standard form."""
    skill_lower = skill.lower().strip()
    for primary, aliases in SKILL_ALIASES.items():
        if skill_lower in aliases:
            return primary
    return skill_lower

def categorize_skill(skill):
    """Determine which category a skill belongs to."""
    skill_lower = skill.lower()
    for category, skills in SKILL_CATEGORIES.items():
        if skill_lower in skills:
            return category
    return 'other'

def calculate_skill_priority(skill, jd_text):
    """Calculate priority of a skill gap based on frequency and context in JD."""
    jd_lower = jd_text.lower()
    skill_lower = skill.lower()
    count = jd_lower.count(skill_lower)

    requirements_keywords = ['required', 'must have', 'essential', 'mandatory', 'minimum']
    in_requirements = any(kw in jd_lower and skill_lower in jd_lower[max(0, jd_lower.find(kw)-200):jd_lower.find(kw)+200]
                          for kw in requirements_keywords)

    preferred_keywords = ['preferred', 'nice to have', 'bonus', 'plus', 'desired']
    in_preferred = any(kw in jd_lower and skill_lower in jd_lower[max(0, jd_lower.find(kw)-200):jd_lower.find(kw)+200]
                       for kw in preferred_keywords)

    if count >= 3 or in_requirements:
        return 'high'
    elif count >= 2 or in_preferred:
        return 'medium'
    else:
        return 'low'

def generate_linkedin_learning_url(skill):
    """Generate a LinkedIn Learning search URL for a skill."""
    encoded_skill = urllib.parse.quote(skill)
    return f"https://www.linkedin.com/learning/search?keywords={encoded_skill}"

def generate_bullet_suggestion(skill, category):
    """Generate a resume bullet point suggestion for a missing skill."""
    suggestions = {
        'programming': f"Developed and maintained applications using {skill}, improving code quality and reducing bugs by X%",
        'frameworks': f"Built scalable web applications leveraging {skill} framework, serving X users",
        'databases': f"Designed and optimized {skill} database schemas, improving query performance by X%",
        'cloud': f"Deployed and managed cloud infrastructure using {skill}, achieving 99.9% uptime",
        'data_science': f"Applied {skill} techniques to analyze datasets, generating actionable insights that increased revenue by X%",
        'soft_skills': f"Demonstrated strong {skill} abilities by leading cross-functional teams to deliver projects on time",
        'tools': f"Utilized {skill} to streamline workflows and improve team productivity by X%",
        'other': f"Leveraged {skill} expertise to drive project success and meet business objectives"
    }
    return suggestions.get(category, suggestions['other'])

def analyze_skill_gaps(resume_text, jd_text, nlp_model):
    """Analyze skill gaps between resume and job description."""
    try:
        # Extract skills from both texts
        resume_skills = extract_skills(resume_text, nlp_model)
        jd_skills = extract_skills(jd_text, nlp_model)

        # Calculate matches and gaps
        matching_skills = resume_skills & jd_skills
        missing_skills = jd_skills - resume_skills
        extra_skills = resume_skills - jd_skills

        # Calculate coverage score
        total_jd_skills = len(jd_skills)
        coverage_score = (len(matching_skills) / total_jd_skills * 100) if total_jd_skills > 0 else 0

        # Generate gap analysis with priorities
        gap_analysis = []
        for skill in missing_skills:
            priority = calculate_skill_priority(skill, jd_text)
            category = categorize_skill(skill)
            suggestion = generate_bullet_suggestion(skill, category)
            learning_url = generate_linkedin_learning_url(skill)

            gap_analysis.append({
                'skill': skill,
                'priority': priority,
                'category': category.replace('_', ' ').title(),
                'suggestion': suggestion,
                'learning_url': learning_url
            })

        # Sort by priority (high -> medium -> low)
        priority_order = {'high': 0, 'medium': 1, 'low': 2}
        gap_analysis.sort(key=lambda x: priority_order.get(x['priority'], 3))

        return {
            'coverage_score': round(coverage_score, 1),
            'matching_skills': sorted(list(matching_skills)),
            'missing_skills': sorted(list(missing_skills)),
            'extra_skills': sorted(list(extra_skills)),
            'gap_analysis': gap_analysis
        }

    except Exception as e:
        print(f"Skill gap analysis failed: {e}")
        return {
            'coverage_score': 0.0,
            'matching_skills': [],
            'missing_skills': [],
            'extra_skills': [],
            'gap_analysis': []
        }

def extract_resume_text(pdf_file):
    """Extract text from PDF resume file."""
    try:
        with pdfplumber.open(pdf_file) as pdf:
            text = ""
            for page in pdf.pages:
                text += page.extract_text() + "\n"
            return text.strip()
    except Exception as e:
        print(f"Error extracting text from PDF: {e}")
        return ""

def fetch_job_description(url):
    """Fetch job description from URL."""
    try:
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        }
        response = requests.get(url, headers=headers, timeout=10)
        response.raise_for_status()

        soup = BeautifulSoup(response.content, 'html.parser')

        # Try to find job description in common selectors
        selectors = [
            '[data-testid="job-description"]',
            '.job-description',
            '.jobsearch-jobDescriptionText',
            '[class*="description"]',
            'article',
            'main'
        ]

        for selector in selectors:
            element = soup.select_one(selector)
            if element:
                text = element.get_text(separator='\n', strip=True)
                if len(text) > 200:  # Must be substantial content
                    return text

        # Fallback: get all paragraph text
        paragraphs = soup.find_all('p')
        text = '\n'.join([p.get_text(strip=True) for p in paragraphs if len(p.get_text(strip=True)) > 20])
        return text if text else "Could not extract job description from URL."

    except Exception as e:
        print(f"Error fetching job description: {e}")
        return f"Error fetching job description: {str(e)}"

def extract_keywords(text, top_n=25):
    """Extract keywords from job description using NLP and frequency analysis."""
    try:
        if not nlp:
            # Fallback without spaCy
            words = text.lower().split()
            # Filter out common stop words
            stop_words = {'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'of', 'with', 'by', 'is', 'are', 'was', 'were', 'be', 'been', 'being', 'have', 'has', 'had', 'do', 'does', 'did', 'will', 'would', 'could', 'should', 'may', 'might', 'must', 'can', 'this', 'that', 'these', 'those', 'i', 'you', 'he', 'she', 'it', 'we', 'they', 'me', 'him', 'her', 'us', 'them'}
            filtered_words = [word for word in words if len(word) > 2 and word not in stop_words]
            word_counts = Counter(filtered_words)
            return [word for word, count in word_counts.most_common(top_n)]

        doc = nlp(text.lower())

        # Extract noun phrases and important nouns
        keywords = []

        # Get noun chunks (phrases)
        for chunk in doc.noun_chunks:
            if len(chunk.text.split()) <= 3:  # Limit to 1-3 word phrases
                keywords.append(chunk.text.strip())

        # Get individual nouns and proper nouns
        for token in doc:
            if token.pos_ in ['NOUN', 'PROPN'] and len(token.text) > 2:
                keywords.append(token.lemma_)

        # Count frequencies
        keyword_counts = Counter(keywords)

        # Filter out overly common words
        filtered_keywords = []
        for keyword, count in keyword_counts.most_common(top_n * 2):
            if count >= 2:  # Must appear at least twice
                filtered_keywords.append(keyword)

        return filtered_keywords[:top_n]

    except Exception as e:
        print(f"Keyword extraction failed: {e}")
        # Simple fallback
        words = text.lower().split()
        word_counts = Counter(words)
        return [word for word, count in word_counts.most_common(top_n) if len(word) > 2]

def calculate_match_score(text1, text2):
    """Calculate semantic similarity between two texts."""
    global EMBEDDING_MODEL
    if not EMBEDDING_MODEL:
        try:
            EMBEDDING_MODEL, _ = load_models()
        except:
            # Fallback to word overlap
            words1 = set(text1.lower().split())
            words2 = set(text2.lower().split())
            common_words = words1 & words2
            if len(words1) > 0 and len(words2) > 0:
                overlap_score = (len(common_words) / max(len(words1), len(words2))) * 100
                return round(overlap_score, 1)
            return 0.0

    try:
        # Use numpy arrays instead of tensors for compatibility
        emb1 = EMBEDDING_MODEL.encode(text1, convert_to_tensor=False)
        emb2 = EMBEDDING_MODEL.encode(text2, convert_to_tensor=False)

        # Calculate cosine similarity
        similarity = util.cos_sim([emb1], [emb2]).item()
        return round(similarity * 100, 1)

    except Exception as e:
        print(f"Embedding calculation failed: {e}")
        # Fallback to simple text overlap if embedding fails
        try:
            words1 = set(text1.lower().split())
            words2 = set(text2.lower().split())
            common_words = words1 & words2

            if len(words1) > 0 and len(words2) > 0:
                overlap_score = (len(common_words) / max(len(words1), len(words2))) * 100
                return round(overlap_score, 1)
            else:
                return 0.0

        except Exception as fallback_e:
            print(f"Fallback method also failed: {str(fallback_e)}")
            return 0.0

def generate_tailored_resume(resume_text, job_text, keywords, llm_config=None):
    """Generate a tailored resume using OpenAI-compatible API."""
    try:
        # Get configured OpenAI client
        client = get_llm_client(llm_config)

        # Determine model name from config or environment
        model = llm_config.get('model') if llm_config else None
        if not model:
            model = os.getenv("LLM_MODEL", "local-model")

        prompt = f"""You are an expert ATS-optimized resume writer with 15+ years of experience helping candidates land interviews at top tech companies. Your task is to strategically rewrite this resume to maximize its compatibility with the target job description while maintaining authenticity and professionalism.

ROLE: Transform the original resume into a compelling, ATS-friendly document that demonstrates clear fit for this specific position through strategic keyword integration, achievement quantification, and relevance optimization.

ORIGINAL RESUME:
{resume_text}

TARGET JOB DESCRIPTION:
{job_text}

KEY REQUIREMENTS IDENTIFIED: {', '.join(keywords[:10])}

OPTIMIZATION FRAMEWORK:

1. KEYWORD INTEGRATION (CRITICAL for ATS):
   - Seamlessly incorporate job-specific keywords naturally into context
   - Use exact keyword phrases from job description where they fit authentically
   - Distribute keywords throughout relevant sections (not just skills list)
   - Prioritize: {', '.join(keywords[:5])} - these must appear prominently

2. ACHIEVEMENT QUANTIFICATION:
   - Convert vague statements into measurable results
   - Add specific metrics, percentages, dollar amounts, or scale indicators
   - Transform "Managed team" → "Managed cross-functional team of 8 engineers"
   - Transform "Improved performance" → "Improved system performance by 40%"

3. EXPERIENCE RELEVANCE:
   - Reorder bullet points to prioritize job-relevant experience first
   - Rephrase past roles to emphasize transferable skills
   - Connect experience directly to job requirements
   - Remove or minimize irrelevant experience

4. SKILLS OPTIMIZATION:
   - Move most relevant skills to the top of skills section
   - Add context around technical skills showing proficiency level
   - Include tools/frameworks mentioned in job description

5. RESUME STRUCTURE & ORGANIZATION:
   - **PRESERVE EXISTING SECTIONS**: Maintain all original sections (Professional Summary, Work Experience, Education, Skills, Certifications, etc.)
   - **STANDARD SECTION HEADERS**: Use consistent, ATS-friendly headers:
     * PROFESSIONAL SUMMARY (or SUMMARY)
     * WORK EXPERIENCE (or EXPERIENCE)
     * EDUCATION
     * SKILLS (or TECHNICAL SKILLS)
     * CERTIFICATIONS (if applicable)
     * PROJECTS (if space allows)
   - **SECTION ORDER**: Professional Summary → Work Experience → Skills → Education → Certifications
   - **CONTENT SEGREGATION**: Never mix different types of content (e.g., don't put education details in work experience section)
   - **BULLET POINT ORGANIZATION**: Group related achievements under appropriate job titles/sections

6. FORMATTING & READABILITY:
   - Maintain clean, scannable format with consistent bullet points
   - Use strong action verbs at start of each bullet
   - Keep similar length but maximize impact per line
   - Ensure professional, error-free language
   - Use consistent date formatting (MM/YYYY)
   - Include location information for remote/hybrid considerations

7. ATS COMPATIBILITY & KEYWORD OPTIMIZATION:
   - Use standard section headers that ATS systems recognize
   - Avoid tables, graphics, or complex formatting
   - Include company names, dates, and location information
   - Distribute keywords naturally throughout relevant sections (not just skills list)
   - Prioritize job-specific keywords: {', '.join(keywords[:8])}
   - Use exact keyword phrases from job description where contextually appropriate

OUTPUT REQUIREMENTS:
- Preserve factual accuracy - do not fabricate experience
- Maintain professional tone appropriate for the industry
- Ensure the resume demonstrates clear career progression
- Make the candidate appear as the ideal fit for this specific role
- Result should be more compelling and targeted than the original

TAILORED RESUME:"""

        # Use OpenAI-compatible API for all providers
        response = client.chat.completions.create(
            model=model,
            messages=[{"role": "user", "content": prompt}],
            max_tokens=2000,
            temperature=0.3
        )

        tailored_resume = response.choices[0].message.content.strip()

        # Clean up any markdown formatting
        tailored_resume = tailored_resume.replace('**', '').replace('*', '')

        return tailored_resume

    except Exception as e:
        print(f"Resume tailoring failed: {e}")
        return resume_text  # Return original if tailoring fails

def create_professional_docx(resume_text, filename="Tailored_Resume.docx"):
    """Create a professional DOCX resume file."""
    try:
        doc = Document()

        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

        # Split resume into lines and process
        lines = resume_text.split('\n')
        current_paragraph = None

        for line in lines:
            line = line.strip()
            if not line:
                continue

            # Check if this is a section header (ALL CAPS or ends with colon)
            if line.isupper() or line.endswith(':'):
                # Add section header
                if current_paragraph:
                    current_paragraph.paragraph_format.space_after = Pt(6)
                p = doc.add_paragraph()
                run = p.add_run(line.upper() if not line.isupper() else line)
                run.font.size = Pt(12)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)
                current_paragraph = p
            else:
                # Regular text
                if current_paragraph:
                    current_paragraph.add_run('\n' + line)
                else:
                    current_paragraph = doc.add_paragraph()
                    current_paragraph.add_run(line)
                current_paragraph.paragraph_format.space_after = Pt(3)

        # Save file
        output_path = os.path.join("outputs", filename)
        os.makedirs("outputs", exist_ok=True)
        doc.save(output_path)
        return output_path

    except Exception as e:
        print(f"Error creating resume DOCX: {e}")
        return None

# ------------------- RECOMMENDATION ANALYSIS -------------------
def generate_strategic_recommendations(resume_text, job_text, skill_analysis, llm_config=None):
    """Generate comprehensive strategic analysis and career recommendations using LLM."""
    try:
        if not llm_config:
            llm_config = {
                'type': 'local',
                'model': 'llama2',
                'server': 'http://localhost:11434'
            }

        client = get_llm_client(llm_config)

        # Prepare comprehensive context
        context = f"""
        RESUME CONTENT:
        {resume_text[:2000]}

        JOB DESCRIPTION:
        {job_text[:2000]}

        SKILL ANALYSIS:
        - Matching Skills: {', '.join(skill_analysis.get('matching_skills', []))}
        - Missing Skills: {', '.join(skill_analysis.get('missing_skills', []))}
        - Coverage Score: {skill_analysis.get('coverage_score', 0)}%
        """

        # Strategic analysis prompt - Enhanced for comprehensive 2-paragraph analysis
        strategic_prompt = f"""
        Based on the resume and job description provided, perform a comprehensive strategic analysis of the candidate's overall fit for this position.

        Context:
        {context}

        Provide a detailed 2-paragraph analysis that covers:

        PARAGRAPH 1: Overall Candidate Assessment
        - Evaluate the candidate's current career level and experience trajectory
        - Assess their demonstrated expertise beyond just technical skills (leadership, business understanding, industry knowledge)
        - Identify their strongest qualifications and unique value propositions
        - Determine if they appear over-qualified, under-qualified, or well-matched for this role

        PARAGRAPH 2: Critical Gaps and Strategic Recommendations
        - Identify the most important missing elements that could prevent hiring (not just skills, but experience gaps, knowledge areas, or intangible factors)
        - Analyze what the employer is really looking for in terms of business impact and strategic contribution
        - Provide specific, actionable recommendations for addressing the most critical gaps
        - Suggest how the candidate can position themselves more competitively

        Then provide specific analysis in these key areas:
        1. Leadership & Management Capability - Evidence of team leadership, project management, or strategic oversight
        2. Industry/Sector Expertise - Deep understanding of this specific industry, market dynamics, and competitive landscape
        3. Business Acumen - Strategic thinking, commercial awareness, and business impact understanding
        4. Technical Depth vs Breadth - Balance of specialized expertise vs adaptability
        5. Communication & Influence - Ability to communicate complex ideas and influence stakeholders
        6. Career Progression & Growth - Pattern of increasing responsibility and achievement
        7. Organizational Fit - Work style, values alignment, and company culture compatibility

        For each area, provide:
        - Priority level (High/Medium/Low) for addressing this gap
        - Detailed analysis explaining the gap's impact on hiring potential
        - Specific recommendations for improvement

        Format as JSON with structure:
        {{
            "strategic_analysis": [
                {{
                    "aspect": "Leadership & Management Capability",
                    "priority": "High/Medium/Low",
                    "analysis": "Detailed 2-3 sentence analysis of this aspect and its importance..."
                }}
            ],
            "career_focus": [
                "Specific, actionable recommendations for career development and positioning..."
            ]
        }}
        """

        response = client.chat.completions.create(
            model=llm_config.get('model', 'local-model'),
            messages=[{"role": "user", "content": strategic_prompt}],
            max_tokens=1500,
            temperature=0.7
        )

        response_text = response.choices[0].message.content.strip()

        # Try to parse JSON response
        try:
            import json
            parsed_response = json.loads(response_text)
            return parsed_response
        except json.JSONDecodeError:
            # Fallback: extract information from text response
            print(f"Failed to parse LLM response as JSON: {response_text[:200]}...")
            return {
                "strategic_analysis": [
                    {
                        "aspect": "Leadership & Management Capability",
                        "priority": "High",
                        "analysis": "The candidate shows limited evidence of leadership experience or team management skills. While they may have individual contributor experience, there's little indication of overseeing projects, managing budgets, or leading cross-functional teams - all critical for senior roles. This represents a significant gap that employers prioritize highly."
                    },
                    {
                        "aspect": "Industry/Sector Expertise",
                        "priority": "Medium",
                        "analysis": "The candidate demonstrates basic familiarity with industry concepts but lacks deep, nuanced understanding of market dynamics, regulatory environments, and competitive landscape that comes from extended experience in this sector. Employers seek candidates who understand industry-specific challenges and opportunities."
                    },
                    {
                        "aspect": "Business Acumen",
                        "priority": "High",
                        "analysis": "There's minimal evidence of strategic thinking, commercial awareness, or understanding of business impact. The resume focuses heavily on technical execution without demonstrating how work contributed to business outcomes, revenue growth, or strategic objectives - a critical gap for roles beyond pure technical implementation."
                    }
                ],
                "career_focus": [
                    "Seek leadership roles or projects that involve team management and decision-making",
                    "Pursue industry-specific certifications and networking opportunities",
                    "Focus on quantifying business impact in resume achievements",
                    "Consider executive education programs for business strategy fundamentals",
                    "Network with senior leaders to understand strategic business challenges"
                ]
            }

    except Exception as e:
        print(f"Error generating strategic recommendations: {e}")
        return {
            "strategic_analysis": [
                {
                    "aspect": "System Analysis Unavailable",
                    "priority": "High",
                    "analysis": f"AI-powered strategic analysis is currently unavailable due to: {str(e)}. However, based on standard industry analysis, most candidates benefit from strengthening leadership experience, industry knowledge, and business acumen beyond technical skills."
                },
                {
                    "aspect": "Standard Career Gaps",
                    "priority": "Medium",
                    "analysis": "Common gaps preventing hiring include lack of demonstrated leadership impact, insufficient industry-specific experience, and limited evidence of strategic thinking. Employers increasingly prioritize candidates who can contribute to business outcomes, not just technical execution."
                }
            ],
            "career_focus": [
                "Focus on gaining leadership experience through team projects or mentoring",
                "Build deeper industry knowledge through targeted networking and certifications",
                "Develop business acumen by understanding company financials and strategic goals",
                "Practice articulating business impact of technical work",
                "Seek roles that involve cross-functional collaboration and decision-making"
            ]
        }

# ------------------- CHAT FUNCTIONALITY -------------------
def process_chat_message(message, context, llm_config=None):
    """Process chat messages with context awareness."""
    try:
        if not llm_config:
            llm_config = {
                'type': 'local',
                'model': 'llama2',
                'server': 'http://localhost:11434'
            }

        client = get_llm_client(llm_config)

        # Build context string
        context_str = ""
        if context:
            context_str = f"""
            JOB DESCRIPTION:
            {context.get('job_description', 'Not available')[:1000]}

            RESUME CONTENT:
            {context.get('resume_content', 'Not available')[:1000]}

            SKILL ANALYSIS:
            Matching: {', '.join(context.get('skill_analysis', {}).get('matching_skills', []))}
            Missing: {', '.join(context.get('skill_analysis', {}).get('missing_skills', []))}
            """

        system_prompt = f"""You are an expert career counselor and resume optimization specialist.
        Help the user with their job search, resume improvement, career development, and interview preparation.

        Context (if available):
        {context_str}

        Guidelines:
        - Be helpful, professional, and encouraging
        - Provide specific, actionable advice
        - Focus on career development and job search strategies
        - Use the provided context when relevant
        - If context is available, reference specific skills or job requirements
        - Be concise but comprehensive
        - Ask clarifying questions when needed

        Current conversation context: The user is working on optimizing their resume for a specific job opportunity.
        """

        messages = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": message}
        ]

        response = client.chat.completions.create(
            model=llm_config.get('model', 'local-model'),
            messages=messages,
            max_tokens=800,
            temperature=0.7
        )

        return response.choices[0].message.content.strip()

    except Exception as e:
        print(f"Error processing chat message: {e}")
        return f"I apologize, but I encountered an error processing your message: {str(e)}. Please check your LLM connection and try again."

# Global variables for model management (initialized lazily)
EMBEDDING_MODEL = None
nlp = None
models_loaded = False
