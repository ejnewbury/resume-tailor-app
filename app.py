# app.py - Resume Tailoring Application

import streamlit as st
import os
import pdfplumber
import requests
from bs4 import BeautifulSoup
import spacy
from collections import Counter
from sentence_transformers import SentenceTransformer, util
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
from openai import OpenAI
import urllib.parse
from jinja2 import Template
import torch
import pandas as pd

# ------------------- PAGE CONFIG & STYLING -------------------
st.set_page_config(
    page_title="Resume Tailoring System",
    page_icon="◈",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Dark minimal UI - Batman computer aesthetic
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=JetBrains+Mono:wght@300;400;500;600&family=Inter:wght@300;400;500;600&display=swap');
    
    .stApp {
        background: #0a0a0a;
        font-family: 'Inter', sans-serif;
    }
    
    .main-header {
        color: #ffffff;
        font-family: 'JetBrains Mono', monospace;
        font-size: 2rem;
        font-weight: 600;
        text-align: center;
        margin-bottom: 0.25rem;
        letter-spacing: 0.1em;
        text-transform: uppercase;
    }
    
    .sub-header {
        color: #404040;
        text-align: center;
        font-size: 0.85rem;
        margin-bottom: 2rem;
        font-family: 'JetBrains Mono', monospace;
        letter-spacing: 0.05em;
    }
    
    .stButton > button {
        background: #1a1a1a;
        color: #00ff88;
        border: 1px solid #00ff88;
        border-radius: 2px;
        padding: 0.75rem 2rem;
        font-weight: 500;
        font-size: 0.85rem;
        font-family: 'JetBrains Mono', monospace;
        text-transform: uppercase;
        letter-spacing: 0.1em;
        transition: all 0.2s ease;
        width: 100%;
    }
    
    .stButton > button:hover {
        background: #00ff88;
        color: #0a0a0a;
        box-shadow: 0 0 20px rgba(0, 255, 136, 0.3);
    }
    
    .sidebar .stButton > button {
        border-color: #00aaff;
        color: #00aaff;
    }
    
    .sidebar .stButton > button:hover {
        background: #00aaff;
        color: #0a0a0a;
        box-shadow: 0 0 20px rgba(0, 170, 255, 0.3);
    }
    
    .stTextInput > div > div > input,
    .stTextArea > div > div > textarea {
        background-color: #0f0f0f;
        border: 1px solid #1f1f1f;
        border-radius: 2px;
        color: #cccccc;
        font-family: 'JetBrains Mono', monospace;
        font-size: 0.85rem;
    }
    
    .stTextInput > div > div > input:focus,
    .stTextArea > div > div > textarea:focus {
        border-color: #00ff88;
        box-shadow: 0 0 10px rgba(0, 255, 136, 0.1);
    }
    
    .stSelectbox > div > div {
        background-color: #0f0f0f;
        border-radius: 2px;
    }
    
    .stRadio > div {
        background: transparent;
    }
    
    .stRadio > div > label {
        color: #666666;
    }
    
    .score-container {
        display: flex;
        justify-content: center;
        align-items: center;
        gap: 2rem;
        margin: 2rem 0;
    }
    
    .score-box {
        background: #0f0f0f;
        border: 1px solid #1a1a1a;
        border-radius: 2px;
        padding: 1.5rem 2rem;
        text-align: center;
        min-width: 180px;
    }
    
    .score-label {
        color: #404040;
        font-family: 'JetBrains Mono', monospace;
        font-size: 0.65rem;
        text-transform: uppercase;
        letter-spacing: 0.15em;
        margin-bottom: 0.5rem;
    }
    
    .score-value {
        font-family: 'JetBrains Mono', monospace;
        font-size: 1.5rem;
        font-weight: 600;
        color: #00ff88;
    }
    
    .score-value.original {
        color: #666666;
    }
    
    .score-value.improvement {
        color: #00aaff;
    }
    
    .preview-box {
        background-color: #0a0a0a;
        border: 1px solid #1a1a1a;
        border-radius: 2px;
        padding: 1.5rem;
        max-height: 500px;
        overflow-y: auto;
        color: #888888;
        font-size: 0.8rem;
        line-height: 1.7;
        font-family: 'JetBrains Mono', monospace;
    }
    
    div[data-testid="stSidebar"] {
        background: #050505;
        border-right: 1px solid #1a1a1a;
    }
    
    div[data-testid="stSidebar"] * {
        color: #888888;
    }
    
    .stAlert {
        background: #0f0f0f;
        border-radius: 2px;
        border: 1px solid #1a1a1a;
    }
    
    h1, h2, h3, h4 {
        color: #ffffff;
        font-family: 'JetBrains Mono', monospace;
    }
    
    .section-header {
        color: #404040;
        font-family: 'JetBrains Mono', monospace;
        font-size: 0.7rem;
        text-transform: uppercase;
        letter-spacing: 0.2em;
        margin-bottom: 1rem;
        padding-bottom: 0.5rem;
        border-bottom: 1px solid #1a1a1a;
    }
    
    .stDownloadButton > button {
        background: #1a1a1a;
        border: 1px solid #00aaff;
        color: #00aaff;
        border-radius: 2px;
    }
    
    .stDownloadButton > button:hover {
        background: #00aaff;
        color: #0a0a0a;
    }
    
    .status-connected {
        color: #00ff88;
        font-family: 'JetBrains Mono', monospace;
        font-size: 0.75rem;
    }
    
    .status-disconnected {
        color: #ff4444;
        font-family: 'JetBrains Mono', monospace;
        font-size: 0.75rem;
    }
    
    .stExpander {
        background: #0a0a0a;
        border: 1px solid #1a1a1a;
        border-radius: 2px;
    }
    
    .stTabs [data-baseweb="tab-list"] {
        background: transparent;
        gap: 0;
    }
    
    .stTabs [data-baseweb="tab"] {
        background: #0a0a0a;
        border: 1px solid #1a1a1a;
        border-radius: 0;
        color: #666666;
        font-family: 'JetBrains Mono', monospace;
        font-size: 0.75rem;
        text-transform: uppercase;
        letter-spacing: 0.1em;
    }
    
    .stTabs [aria-selected="true"] {
        background: #1a1a1a;
        color: #00ff88;
        border-color: #00ff88;
    }
    
    hr {
        border-color: #1a1a1a;
    }
    
    .stFileUploader {
        background: #0a0a0a;
    }
    
    .stFileUploader > div > div {
        background: #0f0f0f;
        border: 1px dashed #1f1f1f;
    }
    
    /* Skill gap analyzer styles */
    .skill-tag {
        display: inline-block;
        padding: 0.25rem 0.5rem;
        margin: 0.15rem;
        border-radius: 2px;
        font-family: 'JetBrains Mono', monospace;
        font-size: 0.7rem;
        text-transform: lowercase;
    }
    
    .skill-match {
        background: rgba(0, 255, 136, 0.15);
        border: 1px solid #00ff88;
        color: #00ff88;
    }
    
    .skill-gap {
        background: rgba(255, 68, 68, 0.15);
        border: 1px solid #ff4444;
        color: #ff4444;
    }
    
    .skill-extra {
        background: rgba(0, 170, 255, 0.15);
        border: 1px solid #00aaff;
        color: #00aaff;
    }
    
    .gap-table {
        width: 100%;
        border-collapse: collapse;
        font-family: 'JetBrains Mono', monospace;
        font-size: 0.75rem;
    }
    
    .gap-table th {
        background: #0f0f0f;
        color: #666666;
        text-transform: uppercase;
        letter-spacing: 0.1em;
        padding: 0.75rem;
        text-align: left;
        border-bottom: 1px solid #1a1a1a;
    }
    
    .gap-table td {
        padding: 0.75rem;
        border-bottom: 1px solid #1a1a1a;
        color: #888888;
    }
    
    .gap-table tr:hover {
        background: #0f0f0f;
    }
    
    .priority-high {
        color: #ff4444;
    }
    
    .priority-medium {
        color: #ffaa00;
    }
    
    .priority-low {
        color: #00aaff;
    }
    
    .learning-link {
        color: #00aaff;
        text-decoration: none;
    }
    
    .learning-link:hover {
        color: #00ff88;
        text-decoration: underline;
    }
    
    .analysis-card {
        background: #0f0f0f;
        border: 1px solid #1a1a1a;
        border-radius: 2px;
        padding: 1rem;
        margin: 0.5rem 0;
    }
    
    .stat-row {
        display: flex;
        justify-content: space-between;
        padding: 0.5rem 0;
        border-bottom: 1px solid #1a1a1a;
    }
    
    .stat-label {
        color: #666666;
        font-family: 'JetBrains Mono', monospace;
        font-size: 0.7rem;
        text-transform: uppercase;
    }
    
    .stat-value {
        color: #ffffff;
        font-family: 'JetBrains Mono', monospace;
        font-size: 0.8rem;
    }
    
    /* Cover letter styles */
    .cover-letter-section {
        background: #0a0a0a;
        border: 1px solid #1a1a1a;
        border-radius: 2px;
        padding: 1.5rem;
        margin-top: 1rem;
    }
    
    .tone-selector {
        display: flex;
        gap: 1rem;
        margin: 1rem 0;
    }
    
    .optional-badge {
        display: inline-block;
        background: rgba(0, 170, 255, 0.15);
        border: 1px solid #00aaff;
        color: #00aaff;
        padding: 0.2rem 0.5rem;
        font-size: 0.6rem;
        text-transform: uppercase;
        letter-spacing: 0.1em;
        margin-left: 0.5rem;
        font-family: 'JetBrains Mono', monospace;
    }
    
    /* Hide streamlit branding */
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
</style>
""", unsafe_allow_html=True)

os.makedirs("outputs", exist_ok=True)

# ------------------- LOAD MODELS -------------------
@st.cache_resource
def load_models(device_override=None):
    """Load models with device selection and error handling."""
    try:
        # Use override if provided, otherwise use session state (default to CPU)
        device = device_override if device_override else getattr(st.session_state, 'device', 'cpu')

        # Load sentence transformer
        embedding_model = SentenceTransformer('all-MiniLM-L6-v2', device=device)

        # Test the model with a small input to catch issues early
        test_text = "test"
        embedding_model.encode(test_text, convert_to_tensor=False)

        # Load spaCy
        nlp_model = spacy.load("en_core_web_sm")

        return embedding_model, nlp_model

    except Exception as e:
        error_msg = str(e).lower()
        if 'cuda' in error_msg or 'gpu' in error_msg or 'kernel' in error_msg:
            st.error("CUDA/GPU error detected. Switching to CPU mode...")
            try:
                # Force CPU mode
                if hasattr(st.session_state, 'device'):
                    st.session_state.device = 'cpu'
                embedding_model = SentenceTransformer('all-MiniLM-L6-v2', device='cpu')
                nlp_model = spacy.load("en_core_web_sm")

                st.warning("Using CPU mode - performance may be slower but more stable")
                return embedding_model, nlp_model

            except Exception as e2:
                st.error(f"Critical error: Could not load models: {str(e2)}")
                st.stop()
        else:
            st.error(f"Model loading error: {str(e)}")
            st.stop()

# Global variables for model management
EMBEDDING_MODEL = None
nlp = None
models_loaded = False

# Load initial models
EMBEDDING_MODEL, nlp = load_models()
models_loaded = True

# ------------------- COVER LETTER TEMPLATES -------------------

COVER_LETTER_TEMPLATES = {
    'professional': Template("""
{{ opening }}

{{ body_intro }}

{{ achievements }}

{{ skills_match }}

{{ closing }}

Sincerely,
{{ candidate_name }}
"""),
    
    'creative': Template("""
{{ opening }}

{{ body_intro }}

{{ achievements }}

{{ skills_match }}

{{ closing }}

Best regards,
{{ candidate_name }}
"""),
    
    'concise': Template("""
{{ opening }}

{{ body_paragraph }}

{{ closing }}

Regards,
{{ candidate_name }}
""")
}

TONE_PROMPTS = {
    'professional': """Write in a formal, professional tone. Use industry-standard language and maintain a respectful, business-appropriate style. 
Focus on qualifications, experience, and demonstrated results. Avoid casual language or overly creative expressions.
Structure: Clear opening stating the position, body paragraphs highlighting relevant experience with metrics, and a confident closing.""",
    
    'creative': """Write in an engaging, memorable tone that shows personality while remaining professional. 
Use compelling storytelling to highlight achievements. Include a unique hook in the opening that captures attention.
Show enthusiasm and passion for the role while demonstrating relevant skills. Be distinctive but not unprofessional.""",
    
    'concise': """Write in a direct, efficient tone. Every sentence must add value - no fluff or filler.
Keep the entire letter under 200 words. Focus only on the 2-3 most relevant qualifications.
Use short paragraphs and punchy sentences. Get straight to the point in the opening."""
}

def extract_candidate_name(resume_text):
    """Extract candidate name from resume (usually first line or prominent text)."""
    lines = resume_text.strip().split('\n')
    for line in lines[:5]:
        line = line.strip()
        # Skip common headers and empty lines
        if line and len(line) < 50 and not any(kw in line.lower() for kw in ['resume', 'cv', 'curriculum', 'email', 'phone', '@', 'http']):
            # Check if it looks like a name (mostly letters, possibly with spaces)
            if sum(c.isalpha() or c.isspace() for c in line) / max(len(line), 1) > 0.8:
                return line.title()
    return "[Your Name]"

def extract_company_info(jd_text):
    """Extract company name and job title from job description."""
    lines = jd_text.strip().split('\n')
    job_title = "[Position Title]"
    company_name = "[Company Name]"
    
    # Common patterns for job titles
    title_keywords = ['position', 'role', 'title', 'job', 'seeking', 'hiring']
    company_keywords = ['company', 'about us', 'who we are', 'organization']
    
    for i, line in enumerate(lines[:20]):
        line_lower = line.lower().strip()
        
        # First non-empty substantive line is often the job title
        if i < 5 and line.strip() and len(line) < 80:
            if not any(kw in line_lower for kw in ['location', 'remote', 'salary', 'posted', 'apply']):
                job_title = line.strip()
                break
    
    # Look for company name patterns
    for line in lines[:30]:
        line_lower = line.lower()
        if 'at ' in line_lower or 'with ' in line_lower or 'join ' in line_lower:
            # Extract potential company name after these keywords
            for keyword in ['at ', 'with ', 'join ']:
                if keyword in line_lower:
                    idx = line_lower.find(keyword) + len(keyword)
                    potential_company = line[idx:idx+50].split(',')[0].split('.')[0].strip()
                    if potential_company and len(potential_company) < 40:
                        company_name = potential_company
                        break
    
    return job_title, company_name

def generate_cover_letter(resume_text, tailored_resume, jd_text, tone, keywords):
    """Generate a cover letter using the connected LLM."""
    
    candidate_name = extract_candidate_name(resume_text)
    job_title, company_name = extract_company_info(jd_text)
    
    tone_instruction = TONE_PROMPTS.get(tone, TONE_PROMPTS['professional'])
    
    prompt = f"""You are an expert cover letter writer. Generate a compelling cover letter for a job application.

TONE REQUIREMENTS:
{tone_instruction}

CANDIDATE NAME: {candidate_name}
TARGET POSITION: {job_title}
TARGET COMPANY: {company_name}

KEY SKILLS TO HIGHLIGHT (from job description):
{', '.join(keywords[:15])}

TAILORED RESUME CONTENT:
{tailored_resume[:6000]}

JOB DESCRIPTION:
{jd_text[:4000]}

REQUIREMENTS:
1. Address the letter appropriately (Dear Hiring Manager, or Dear [Company] Team)
2. Open with a strong hook that shows enthusiasm and states the position
3. Connect 2-3 specific achievements from the resume to job requirements
4. Demonstrate knowledge of the company/role where possible
5. Include a confident call to action in the closing
6. Keep total length appropriate for the tone (professional: 300-400 words, creative: 350-450 words, concise: 150-200 words)
7. Do NOT use placeholder brackets like [Company] - use the actual company name or write generically if unknown
8. Sign off with the candidate's name: {candidate_name}

OUTPUT:
Return ONLY the cover letter text. No explanations, no markdown formatting, no headers like "Cover Letter:".
Start directly with the salutation (Dear...)."""

    if st.session_state.client is None:
        st.error("No LLM connection established")
        return None

    try:
        if hasattr(st.session_state, 'model_choice'):
            model = st.session_state.model_choice
        else:
            model = "local-model"
        
        response = st.session_state.client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": f"You are an expert cover letter writer specializing in {tone} tone letters that get interviews."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.8 if tone == 'creative' else 0.7,
            max_tokens=1500
        )
        
        return response.choices[0].message.content
        
    except Exception as e:
        st.error(f"Generation error: {str(e)}")
        return None

def create_cover_letter_docx(cover_letter_text, filename="Cover_Letter.docx"):
    """Create a professionally formatted cover letter Word document."""
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Add date
    date_para = doc.add_paragraph()
    date_run = date_para.add_run(datetime.now().strftime("%B %d, %Y"))
    date_run.font.size = Pt(11)
    date_para.space_after = Pt(24)
    
    # Add cover letter content
    paragraphs = cover_letter_text.strip().split('\n\n')
    
    for para_text in paragraphs:
        para_text = para_text.strip()
        if not para_text:
            continue
        
        # Handle single line breaks within paragraphs
        para_text = para_text.replace('\n', ' ').strip()
        
        p = doc.add_paragraph(para_text)
        for run in p.runs:
            run.font.size = Pt(11)
            run.font.name = 'Calibri'
        p.paragraph_format.space_after = Pt(12)
        p.paragraph_format.line_spacing = 1.15
    
    filepath = os.path.join("outputs", filename)
    doc.save(filepath)
    return filepath

# ------------------- SKILL EXTRACTION & GAP ANALYSIS -------------------

SKILL_ALIASES = {
    'javascript': ['js', 'javascript', 'ecmascript'],
    'typescript': ['ts', 'typescript'],
    'python': ['python', 'py'],
    'machine learning': ['ml', 'machine learning', 'machine-learning'],
    'artificial intelligence': ['ai', 'artificial intelligence'],
    'amazon web services': ['aws', 'amazon web services'],
    'google cloud platform': ['gcp', 'google cloud', 'google cloud platform'],
    'microsoft azure': ['azure', 'microsoft azure'],
    'continuous integration': ['ci', 'continuous integration'],
    'continuous deployment': ['cd', 'continuous deployment'],
    'ci/cd': ['ci/cd', 'cicd', 'ci cd'],
    'kubernetes': ['k8s', 'kubernetes'],
    'postgresql': ['postgres', 'postgresql'],
    'mongodb': ['mongo', 'mongodb'],
    'react': ['react', 'reactjs', 'react.js'],
    'angular': ['angular', 'angularjs'],
    'vue': ['vue', 'vuejs', 'vue.js'],
    'node.js': ['node', 'nodejs', 'node.js'],
    'sql': ['sql', 'structured query language'],
    'nosql': ['nosql', 'no-sql'],
    'api': ['api', 'apis', 'rest api', 'restful'],
    'agile': ['agile', 'scrum', 'kanban'],
    'project management': ['project management', 'pm', 'program management'],
    'data analysis': ['data analysis', 'data analytics', 'analytics'],
    'data science': ['data science', 'data scientist'],
    'deep learning': ['deep learning', 'dl', 'neural networks'],
    'natural language processing': ['nlp', 'natural language processing'],
    'computer vision': ['cv', 'computer vision', 'image processing'],
}

SKILL_CATEGORIES = {
    'programming': ['python', 'java', 'javascript', 'typescript', 'c++', 'c#', 'ruby', 'go', 'rust', 'php', 'swift', 'kotlin', 'scala', 'r'],
    'frameworks': ['react', 'angular', 'vue', 'django', 'flask', 'spring', 'express', 'rails', 'laravel', 'nextjs', 'fastapi'],
    'databases': ['sql', 'mysql', 'postgresql', 'mongodb', 'redis', 'elasticsearch', 'cassandra', 'dynamodb', 'oracle', 'sqlite'],
    'cloud': ['aws', 'azure', 'gcp', 'docker', 'kubernetes', 'terraform', 'jenkins', 'circleci', 'github actions'],
    'data_science': ['machine learning', 'deep learning', 'tensorflow', 'pytorch', 'pandas', 'numpy', 'scikit-learn', 'spark', 'hadoop'],
    'soft_skills': ['leadership', 'communication', 'teamwork', 'problem solving', 'analytical', 'strategic', 'collaboration', 'mentoring'],
    'tools': ['git', 'jira', 'confluence', 'slack', 'figma', 'photoshop', 'excel', 'tableau', 'power bi'],
}

def extract_skills(text, nlp_model):
    """Extract skills from text using spaCy NLP and pattern matching."""
    doc = nlp_model(text.lower())
    skills = set()
    
    noun_phrases = [chunk.text.strip() for chunk in doc.noun_chunks]
    tokens = [token.lemma_ for token in doc if token.is_alpha and len(token.text) > 1]
    candidates = set(noun_phrases + tokens)
    
    all_known_skills = set()
    for category_skills in SKILL_CATEGORIES.values():
        all_known_skills.update(category_skills)
    for skill, aliases in SKILL_ALIASES.items():
        all_known_skills.add(skill)
        all_known_skills.update(aliases)
    
    for candidate in candidates:
        candidate_clean = candidate.strip().lower()
        if candidate_clean in all_known_skills:
            normalized = normalize_skill(candidate_clean)
            skills.add(normalized)
            continue
        for skill in all_known_skills:
            if skill in candidate_clean or candidate_clean in skill:
                normalized = normalize_skill(skill)
                skills.add(normalized)
    
    text_lower = text.lower()
    skill_patterns = [
        'experience with', 'proficient in', 'knowledge of', 'skilled in',
        'expertise in', 'familiar with', 'working knowledge of', 'hands-on experience',
        'strong background in', 'demonstrated ability'
    ]
    
    for pattern in skill_patterns:
        if pattern in text_lower:
            idx = text_lower.find(pattern)
            snippet = text_lower[idx:idx+100]
            snippet_doc = nlp_model(snippet)
            for token in snippet_doc:
                if token.lemma_ in all_known_skills:
                    skills.add(normalize_skill(token.lemma_))
    
    return skills

def normalize_skill(skill):
    """Normalize skill name to a standard form."""
    skill_lower = skill.lower().strip()
    for primary, aliases in SKILL_ALIASES.items():
        if skill_lower in aliases:
            return primary
    return skill_lower

def categorize_skill(skill):
    """Determine which category a skill belongs to."""
    skill_lower = skill.lower()
    for category, skills in SKILL_CATEGORIES.items():
        if skill_lower in skills:
            return category
    return 'other'

def calculate_skill_priority(skill, jd_text):
    """Calculate priority of a skill gap based on frequency and context in JD."""
    jd_lower = jd_text.lower()
    skill_lower = skill.lower()
    count = jd_lower.count(skill_lower)
    
    requirements_keywords = ['required', 'must have', 'essential', 'mandatory', 'minimum']
    in_requirements = any(kw in jd_lower and skill_lower in jd_lower[max(0, jd_lower.find(kw)-200):jd_lower.find(kw)+200] 
                          for kw in requirements_keywords)
    
    preferred_keywords = ['preferred', 'nice to have', 'bonus', 'plus', 'desired']
    in_preferred = any(kw in jd_lower and skill_lower in jd_lower[max(0, jd_lower.find(kw)-200):jd_lower.find(kw)+200] 
                       for kw in preferred_keywords)
    
    if count >= 3 or in_requirements:
        return 'high'
    elif count >= 2 or in_preferred:
        return 'medium'
    else:
        return 'low'

def generate_linkedin_learning_url(skill):
    """Generate a LinkedIn Learning search URL for a skill."""
    encoded_skill = urllib.parse.quote(skill)
    return f"https://www.linkedin.com/learning/search?keywords={encoded_skill}"

def generate_bullet_suggestion(skill, category):
    """Generate a resume bullet point suggestion for a missing skill."""
    suggestions = {
        'programming': f"Developed and maintained applications using {skill}, improving code quality and reducing bugs by X%",
        'frameworks': f"Built scalable web applications leveraging {skill} framework, serving X users",
        'databases': f"Designed and optimized {skill} database schemas, improving query performance by X%",
        'cloud': f"Deployed and managed cloud infrastructure using {skill}, achieving 99.9% uptime",
        'data_science': f"Applied {skill} techniques to analyze datasets, generating actionable insights that increased revenue by X%",
        'soft_skills': f"Demonstrated strong {skill} abilities by leading cross-functional teams to deliver projects on time",
        'tools': f"Utilized {skill} to streamline workflows and improve team productivity by X%",
        'other': f"Leveraged {skill} expertise to drive project success and meet business objectives"
    }
    return suggestions.get(category, suggestions['other'])

def analyze_skill_gaps(resume_text, jd_text, nlp_model):
    """Comprehensive skill gap analysis between resume and job description."""
    resume_skills = extract_skills(resume_text, nlp_model)
    jd_skills = extract_skills(jd_text, nlp_model)
    
    matching_skills = resume_skills & jd_skills
    missing_skills = jd_skills - resume_skills
    extra_skills = resume_skills - jd_skills
    
    gap_analysis = []
    for skill in missing_skills:
        category = categorize_skill(skill)
        priority = calculate_skill_priority(skill, jd_text)
        gap_analysis.append({
            'skill': skill,
            'category': category,
            'priority': priority,
            'suggestion': generate_bullet_suggestion(skill, category),
            'learning_url': generate_linkedin_learning_url(skill)
        })
    
    priority_order = {'high': 0, 'medium': 1, 'low': 2}
    gap_analysis.sort(key=lambda x: priority_order[x['priority']])
    
    if len(jd_skills) > 0:
        coverage_score = round((len(matching_skills) / len(jd_skills)) * 100, 1)
    else:
        coverage_score = 100.0
    
    return {
        'resume_skills': resume_skills,
        'jd_skills': jd_skills,
        'matching_skills': matching_skills,
        'missing_skills': missing_skills,
        'extra_skills': extra_skills,
        'gap_analysis': gap_analysis,
        'coverage_score': coverage_score
    }

# ------------------- SESSION STATE -------------------
if 'client' not in st.session_state:
    st.session_state.client = None
if 'connected' not in st.session_state:
    st.session_state.connected = False
if 'llm_info' not in st.session_state:
    st.session_state.llm_info = ""
if 'tailored_resume' not in st.session_state:
    st.session_state.tailored_resume = None
if 'cover_letter' not in st.session_state:
    st.session_state.cover_letter = None
if 'device' not in st.session_state:
    st.session_state.device = 'cpu'

# ------------------- SIDEBAR: SYSTEM SETTINGS -------------------
st.sidebar.markdown("### SYSTEM SETTINGS")
st.sidebar.markdown("---")

# Device selection
st.sidebar.markdown('<p class="section-header">Device Configuration</p>', unsafe_allow_html=True)

device_options = ["Auto (GPU if available)", "CPU Only"]
device_choice = st.sidebar.selectbox(
    "Processing Device",
    device_options,
    help="Choose CPU for stability, Auto for best performance",
    key="device_selector"
)

# Handle device changes
previous_device = st.session_state.get('previous_device', None)
current_device_choice = device_choice

if device_choice == "CPU Only":
    new_device = 'cpu'
else:
    new_device = 'cuda' if torch.cuda.is_available() else 'cpu'

# Check if device changed
if previous_device != current_device_choice or st.session_state.device != new_device:
    st.session_state.device = new_device
    st.session_state.previous_device = current_device_choice

    # Clear cache to force reload models with new device
    load_models.clear()

    # Reload models with new device
    try:
        # Force clear cache and reload
        load_models.clear()
        new_embedding_model, new_nlp = load_models(new_device)

        # Update global references
        globals()['EMBEDDING_MODEL'] = new_embedding_model
        globals()['nlp'] = new_nlp
        globals()['models_loaded'] = True

        if new_device == 'cpu':
            st.sidebar.info("Using CPU mode")
        else:
            st.sidebar.success("GPU available and active")

    except Exception as e:
        st.sidebar.error(f"Device switch failed: {str(e)}")
        globals()['models_loaded'] = False

def setup_sidebar_status():
    """Setup all sidebar status and troubleshooting elements that require session state."""
    # Display current device status
    if st.session_state.device == 'cpu':
        st.sidebar.info("● CPU Mode")
    else:
        st.sidebar.success("● GPU Mode")

    # CUDA troubleshooting tips
    with st.sidebar.expander("CUDA Troubleshooting", expanded=False):
        st.markdown("""
        **If you encounter CUDA errors:**
        - Run with: `CUDA_LAUNCH_BLOCKING=1 streamlit run app.py`
        - Or set: `TORCH_USE_CUDA_DSA=1`
        - Check GPU compatibility with your PyTorch version
        - Force CPU mode using the selector above
        """)

    st.sidebar.markdown("---")

    # Connection status
    st.sidebar.markdown('<p class="section-header">Status</p>', unsafe_allow_html=True)
    if st.session_state.connected:
        st.sidebar.markdown(f'<p class="status-connected">● ONLINE: {st.session_state.llm_info}</p>', unsafe_allow_html=True)
    else:
        st.sidebar.markdown('<p class="status-disconnected">● OFFLINE</p>', unsafe_allow_html=True)

# ------------------- SIDEBAR: LLM CONNECTION -------------------
st.sidebar.markdown("### LLM CONNECTION")
st.sidebar.markdown("---")

llm_type = st.sidebar.radio(
    "Connection Type",
    ["Local LLM", "OpenAI Cloud"],
    help="Select LLM connection method"
)

if llm_type == "Local LLM":
    st.sidebar.markdown('<p class="section-header">Local Server Configuration</p>', unsafe_allow_html=True)
    
    model_name = st.sidebar.text_input(
        "Model Identifier",
        placeholder="llama-3, mistral, dolphin...",
        help="Model name for reference"
    )
    
    ip_address = st.sidebar.text_input(
        "Server Address",
        value="http://localhost:1234",
        placeholder="http://localhost:1234",
        help="LLM server endpoint (LM Studio, Ollama, etc.)"
    )
    
    if st.sidebar.button("CONNECT", use_container_width=True):
        if not ip_address:
            st.sidebar.error("Server address required")
        else:
            try:
                base_url = ip_address.rstrip('/')
                if not base_url.endswith('/v1'):
                    base_url = f"{base_url}/v1"

                st.session_state.client = OpenAI(
                    base_url=base_url,
                    api_key="not-needed"
                )

                st.session_state.client.models.list()

                st.session_state.connected = True
                st.session_state.llm_info = f"{model_name or 'Local'} @ {ip_address}"
                st.sidebar.success("Connection established")

            except Exception as e:
                st.session_state.connected = False
                st.sidebar.error(f"Connection failed: {str(e)}")

elif llm_type == "OpenAI Cloud":
    st.sidebar.markdown('<p class="section-header">OpenAI Configuration</p>', unsafe_allow_html=True)
    
    api_key = st.sidebar.text_input(
        "API Key",
        type="password",
        placeholder="sk-...",
        help="OpenAI API key"
    )
    
    model_choice = st.sidebar.selectbox(
        "Model",
        ["gpt-4o-mini", "gpt-4o", "gpt-4-turbo", "gpt-3.5-turbo"]
    )
    
    if st.sidebar.button("CONNECT", use_container_width=True):
        if not api_key:
            st.sidebar.error("API key required")
        elif not api_key.startswith("sk-"):
            st.sidebar.error("Invalid API key format")
        else:
            try:
                st.session_state.client = OpenAI(api_key=api_key)
                st.session_state.client.models.list()
                
                st.session_state.connected = True
                st.session_state.llm_info = f"OpenAI {model_choice}"
                st.session_state.model_choice = model_choice
                st.sidebar.success("Connection established")
                
            except Exception as e:
                st.session_state.connected = False
                st.sidebar.error(f"Connection failed: {str(e)}")

# Setup sidebar status (only when in Streamlit context)
try:
    # Only call this when running in Streamlit
    if hasattr(st, 'session_state'):
        setup_sidebar_status()
except:
    # Skip if not in Streamlit context (e.g., when imported by API)
    pass

# ------------------- HELPER FUNCTIONS -------------------

def extract_resume_text(pdf_file):
    """Extract text from uploaded PDF resume."""
    with pdfplumber.open(pdf_file) as pdf:
        return "\n".join(page.extract_text() or "" for page in pdf.pages)

def fetch_job_description(url):
    """Fetch job description from URL."""
    try:
        headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'}
        response = requests.get(url, headers=headers, timeout=10)
        soup = BeautifulSoup(response.text, 'html.parser')
        
        for tag in soup(['script', 'style', 'nav', 'header', 'footer', 'aside', 'form']):
            tag.decompose()
        
            return soup.get_text(separator='\n', strip=True)
    except Exception as e:
        return None

def extract_keywords(text, top_n=25):
    """Extract important keywords from text using NLP."""
    doc = nlp(text.lower())
    
    words = [
        token.lemma_ for token in doc 
        if token.is_alpha 
        and not token.is_stop 
        and len(token.text) > 2
        and token.pos_ in ['NOUN', 'VERB', 'ADJ', 'PROPN']
    ]
    
    return [word for word, count in Counter(words).most_common(top_n)]

def calculate_match_score(text1, text2):
    """Calculate semantic similarity between two texts with error handling."""
    try:
        # Use numpy arrays instead of tensors for compatibility
        emb1 = EMBEDDING_MODEL.encode(text1, convert_to_tensor=False)
        emb2 = EMBEDDING_MODEL.encode(text2, convert_to_tensor=False)

        # Calculate cosine similarity
        similarity = util.cos_sim([emb1], [emb2]).item()
        return round(similarity * 100, 1)

    except Exception as e:
        # Fallback to simple text overlap if embedding fails
        st.warning(f"Embedding calculation failed: {str(e)}. Using fallback method.")

        try:
            # Simple fallback: count common words
            words1 = set(text1.lower().split())
            words2 = set(text2.lower().split())
            common_words = words1 & words2

            if len(words1) > 0 and len(words2) > 0:
                overlap_score = (len(common_words) / max(len(words1), len(words2))) * 100
                return round(overlap_score, 1)
            else:
                return 0.0

        except Exception as fallback_e:
            st.error(f"Fallback method also failed: {str(fallback_e)}")
            return 0.0

def generate_tailored_resume(resume_text, job_text, keywords, model_name=None):
    """Generate a tailored resume using the connected LLM."""
    
    prompt = f"""You are an expert professional resume writer with 20+ years of experience helping candidates land interviews at top companies.

TASK: Completely rewrite and optimize the provided resume to perfectly match the target job description.

REQUIREMENTS:
1. **ATS Optimization**: Use exact keywords and phrases from the job description naturally throughout
2. **Quantify Achievements**: Add metrics, percentages, and numbers wherever possible (estimate if needed)
3. **Action Verbs**: Start each bullet with strong action verbs (Led, Developed, Implemented, Achieved, etc.)
4. **Relevance**: Prioritize and expand experiences most relevant to the target role
5. **Modern Format**: Use clean, scannable formatting with clear sections
6. **Concise**: Keep bullets impactful - ideally 1-2 lines each
7. **Professional Summary**: Write a compelling 3-4 sentence summary tailored to this specific role

CRITICAL KEYWORDS TO INCORPORATE NATURALLY:
{', '.join(keywords[:20])}

TARGET JOB DESCRIPTION:
{job_text[:5000]}

ORIGINAL RESUME:
{resume_text[:10000]}

OUTPUT FORMAT:
Return the complete rewritten resume in clean, professional text format with clear section headers.
Use this structure:
- PROFESSIONAL SUMMARY
- SKILLS
- PROFESSIONAL EXPERIENCE
- EDUCATION
- (Any other relevant sections)

Do NOT include any explanations, notes, or markdown formatting. Just the resume content."""

    if st.session_state.client is None:
        st.error("No LLM connection established")
        return None

    try:
        if hasattr(st.session_state, 'model_choice'):
            model = st.session_state.model_choice
        else:
            model = model_name or "local-model"
        
        response = st.session_state.client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": "You are an expert resume writer who creates ATS-optimized, professional resumes."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.7,
            max_tokens=4000
        )
        
        return response.choices[0].message.content
        
    except Exception as e:
        st.error(f"Generation error: {str(e)}")
        return None

def create_professional_docx(resume_text, filename="Tailored_Resume.docx"):
    """Create a professionally formatted Word document."""
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    lines = resume_text.strip().split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue
        
        is_header = (
            line.isupper() and len(line) > 3 and len(line) < 50
        ) or any(header in line.upper() for header in [
            'SUMMARY', 'EXPERIENCE', 'EDUCATION', 'SKILLS', 
            'CERTIFICATIONS', 'PROJECTS', 'ACHIEVEMENTS'
        ])
        
        if is_header:
            p = doc.add_paragraph()
            run = p.add_run(line.upper())
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0, 51, 102)
            p.space_after = Pt(6)
        elif line.startswith(('•', '-', '●', '*')):
            p = doc.add_paragraph(line[1:].strip(), style='List Bullet')
            p.paragraph_format.space_after = Pt(3)
            for run in p.runs:
                run.font.size = Pt(10.5)
        else:
            p = doc.add_paragraph(line)
            for run in p.runs:
                run.font.size = Pt(10.5)
            p.paragraph_format.space_after = Pt(3)
    
    filepath = os.path.join("outputs", filename)
    doc.save(filepath)
    return filepath

# ------------------- MAIN APPLICATION -------------------

st.markdown('<h1 class="main-header">Resume Tailoring System</h1>', unsafe_allow_html=True)
st.markdown('<p class="sub-header">// Optimize your resume for any position</p>', unsafe_allow_html=True)

# Main content area
col1, col2 = st.columns([1, 1], gap="large")

with col1:
    st.markdown('<p class="section-header">Input: Resume</p>', unsafe_allow_html=True)
    uploaded_file = st.file_uploader(
        "Upload PDF",
        type="pdf",
        help="Upload your current resume in PDF format",
        label_visibility="collapsed"
    )
    
    if uploaded_file:
        st.success(f"Loaded: {uploaded_file.name}")
        
        with st.expander("Preview Original", expanded=False):
            resume_text = extract_resume_text(uploaded_file)
            st.markdown(f'<div class="preview-box">{resume_text[:2000]}...</div>', unsafe_allow_html=True)

with col2:
    st.markdown('<p class="section-header">Input: Job Description</p>', unsafe_allow_html=True)
    
    input_method = st.tabs(["Paste Text", "From URL"])
    
    with input_method[0]:
        job_text_input = st.text_area(
            "Job Description",
            height=200,
            placeholder="Paste the full job description here...",
            label_visibility="collapsed"
        )
    
    with input_method[1]:
        job_url_input = st.text_input(
            "Job Posting URL",
            placeholder="https://...",
            label_visibility="collapsed"
        )
        if job_url_input:
            with st.spinner("Fetching..."):
                fetched_text = fetch_job_description(job_url_input)
                if fetched_text:
                    st.success("Job description fetched")
                    with st.expander("Preview fetched content", expanded=False):
                        st.text(fetched_text[:1000] + "...")
                else:
                    st.error("Failed to fetch URL")
    
    job_input = job_text_input if job_text_input else (fetch_job_description(job_url_input) if job_url_input else None)
    
    if job_input:
        keywords = extract_keywords(job_input)
        with st.expander("Extracted Keywords", expanded=False):
            st.text(" | ".join(keywords))

# Generate button and results
st.markdown("---")

if uploaded_file and job_input:
    col_btn1, col_btn2, col_btn3 = st.columns([1, 2, 1])
    
    with col_btn2:
        generate_btn = st.button(
            "GENERATE TAILORED RESUME",
            use_container_width=True,
            disabled=not st.session_state.connected
        )
    
    if not st.session_state.connected:
        st.warning("Connect to an LLM in the sidebar to continue")
    
    if generate_btn and st.session_state.connected:
        resume_text = extract_resume_text(uploaded_file)
        keywords = extract_keywords(job_input)
        
        initial_score = calculate_match_score(resume_text, job_input)
        
        with st.spinner("Processing..."):
            tailored_resume = generate_tailored_resume(resume_text, job_input, keywords)
            st.session_state.tailored_resume = tailored_resume
            st.session_state.resume_text = resume_text
            st.session_state.job_input = job_input
            st.session_state.keywords = keywords
        
            if tailored_resume:
                new_score = calculate_match_score(tailored_resume, job_input)
                improvement = new_score - initial_score

                st.success("Generation complete")

                # Score display
                st.markdown(f"""
            <div class="score-container">
                <div class="score-box">
                    <div class="score-label">Original Score</div>
                    <div class="score-value original">{initial_score}%</div>
                </div>
                <div class="score-box">
                    <div class="score-label">Improvement</div>
                    <div class="score-value improvement">+{improvement:.1f}%</div>
                </div>
                <div class="score-box">
                    <div class="score-label">Optimized Score</div>
                    <div class="score-value">{new_score}%</div>
                </div>
            </div>
            """, unsafe_allow_html=True)
            
            st.markdown("---")
            
            # ------------------- SKILL GAP ANALYZER -------------------
            st.markdown('<p class="section-header">Skill Gap Analysis</p>', unsafe_allow_html=True)
            
            with st.spinner("Analyzing skill gaps..."):
                gap_results = analyze_skill_gaps(resume_text, job_input, nlp)
            
            st.markdown(f"""
            <div class="analysis-card">
                <div class="stat-row">
                    <span class="stat-label">Skill Coverage</span>
                    <span class="stat-value">{gap_results['coverage_score']}%</span>
                </div>
                <div class="stat-row">
                    <span class="stat-label">Matching Skills</span>
                    <span class="stat-value">{len(gap_results['matching_skills'])}</span>
                </div>
                <div class="stat-row">
                    <span class="stat-label">Missing Skills</span>
                    <span class="stat-value">{len(gap_results['missing_skills'])}</span>
                </div>
                <div class="stat-row" style="border-bottom: none;">
                    <span class="stat-label">Additional Skills</span>
                    <span class="stat-value">{len(gap_results['extra_skills'])}</span>
                </div>
            </div>
            """, unsafe_allow_html=True)
            
            with st.expander("Skills Breakdown", expanded=True):
                skill_col1, skill_col2 = st.columns(2)
                
                with skill_col1:
                    st.markdown("**Matching Skills**")
                    if gap_results['matching_skills']:
                        skills_html = " ".join([f'<span class="skill-tag skill-match">{skill}</span>' for skill in sorted(gap_results['matching_skills'])])
                        st.markdown(skills_html, unsafe_allow_html=True)
                    else:
                        st.text("No matching skills found")
                
                with skill_col2:
                    st.markdown("**Your Additional Skills**")
                    if gap_results['extra_skills']:
                        skills_html = " ".join([f'<span class="skill-tag skill-extra">{skill}</span>' for skill in sorted(gap_results['extra_skills'])])
                        st.markdown(skills_html, unsafe_allow_html=True)
                    else:
                        st.text("None")
            
            if gap_results['gap_analysis']:
                with st.expander("Missing Skills - Action Plan", expanded=True):
                    st.markdown("**Skills to Add to Your Resume:**")

                    # Create data for the table
                    table_data = []
                    for gap in gap_results['gap_analysis']:
                        table_data.append({
                            'Skill': f"🔴 {gap['skill']}",  # Red circle for missing skills
                            'Priority': gap['priority'].upper(),
                            'Category': gap['category'].replace('_', ' ').title(),
                            'Suggested Bullet': gap['suggestion'][:100] + '...' if len(gap['suggestion']) > 100 else gap['suggestion'],
                            'Learning Link': gap['learning_url']
                        })

                    # Display as simple table
                    st.table(table_data)

                    st.markdown("---")
                    st.markdown("**Priority Legend:**")
                    st.markdown("🔴 **HIGH** = Mentioned 3+ times or in requirements section")
                    st.markdown("🟡 **MEDIUM** = Mentioned 2 times or in preferred section")
                    st.markdown("🔵 **LOW** = Nice to have or mentioned once")

                    # Individual learning links
                    st.markdown("### Quick Learning Links:")
                    cols = st.columns(min(3, len(gap_results['gap_analysis'])))
                    for i, gap in enumerate(gap_results['gap_analysis']):
                        col_idx = i % 3
                        with cols[col_idx]:
                            st.markdown(f"[{gap['skill'].title()}]({gap['learning_url']})")
            
            st.markdown("---")
            
            # Preview and download tailored resume
            st.markdown('<p class="section-header">Output: Tailored Resume</p>', unsafe_allow_html=True)
            st.markdown(f'<div class="preview-box">{tailored_resume}</div>', unsafe_allow_html=True)
            
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"Tailored_Resume_{timestamp}.docx"
            filepath = create_professional_docx(tailored_resume, filename)
            
            st.markdown("---")
            
            dl_col1, dl_col2, dl_col3 = st.columns([1, 2, 1])
            with dl_col2:
                with open(filepath, "rb") as f:
                    st.download_button(
                        "DOWNLOAD RESUME DOCX",
                        data=f,
                        file_name=filename,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
            
            # ------------------- COVER LETTER GENERATOR -------------------
            st.markdown("---")
            st.markdown('<p class="section-header">Cover Letter Generator <span class="optional-badge">Optional</span></p>', unsafe_allow_html=True)
            
            st.markdown("Generate a personalized cover letter that complements your tailored resume.")
            
            # Tone selection
            tone_col1, tone_col2 = st.columns([2, 3])
            
            with tone_col1:
                cover_letter_tone = st.radio(
                    "Select Tone",
                    ["Professional", "Creative", "Concise"],
                    help="Professional: Formal and business-appropriate | Creative: Engaging with personality | Concise: Direct and under 200 words",
                    horizontal=True
                )
            
            with tone_col2:
                tone_descriptions = {
                    "Professional": "Formal, business-appropriate language with focus on qualifications and results.",
                    "Creative": "Engaging storytelling with personality while remaining professional.",
                    "Concise": "Direct and efficient, under 200 words. Every sentence adds value."
                }
                st.info(tone_descriptions[cover_letter_tone])
            
            # Optional personalization inputs
            with st.expander("Personalization Options", expanded=False):
                pers_col1, pers_col2 = st.columns(2)
                with pers_col1:
                    hiring_manager_name = st.text_input(
                        "Hiring Manager Name",
                        placeholder="Leave blank for 'Dear Hiring Manager'",
                        help="If known, enter the hiring manager's name"
                    )
                with pers_col2:
                    company_detail = st.text_input(
                        "Why This Company",
                        placeholder="e.g., 'innovative AI products', 'company culture'",
                        help="Mention something specific about why you want to work there"
                    )
            
            # Generate cover letter button
            cl_btn_col1, cl_btn_col2, cl_btn_col3 = st.columns([1, 2, 1])
            with cl_btn_col2:
                generate_cl_btn = st.button(
                    "GENERATE COVER LETTER",
                    use_container_width=True,
                    key="generate_cover_letter"
                )
            
            if generate_cl_btn:
                with st.spinner("Generating cover letter..."):
                    cover_letter = generate_cover_letter(
                        resume_text=resume_text,
                        tailored_resume=tailored_resume,
                        jd_text=job_input,
                        tone=cover_letter_tone.lower(),
                        keywords=keywords
                    )
                    st.session_state.cover_letter = cover_letter
                
                if cover_letter:
                    st.success("Cover letter generated")
                    
                    st.markdown('<p class="section-header">Output: Cover Letter</p>', unsafe_allow_html=True)
                    st.markdown(f'<div class="preview-box">{cover_letter}</div>', unsafe_allow_html=True)
                    
                    # Create and offer download
                    cl_timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                    cl_filename = f"Cover_Letter_{cl_timestamp}.docx"
                    cl_filepath = create_cover_letter_docx(cover_letter, cl_filename)
                    
                    st.markdown("---")
                    
                    cl_dl_col1, cl_dl_col2, cl_dl_col3 = st.columns([1, 2, 1])
                    with cl_dl_col2:
                        with open(cl_filepath, "rb") as f:
                            st.download_button(
                                "DOWNLOAD COVER LETTER DOCX",
                                data=f,
                                file_name=cl_filename,
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                use_container_width=True,
                                key="download_cover_letter"
                            )

else:
    st.info("Upload a resume and provide a job description to begin")
